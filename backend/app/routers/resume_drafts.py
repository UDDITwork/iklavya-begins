import json
import io
import os
import re
import tempfile
from datetime import datetime, timezone

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, status
from fastapi.responses import Response
from sqlalchemy.orm import Session

from app.database import get_db
from app.models import User, UserProfile, ResumeDraft
from app.schemas import (
    ResumeDraftCreateRequest,
    ResumeDraftResponse,
    ResumeDraftListResponse,
    ResumeDraftUpdateRequest,
    ResumeDraftAIOptimizeRequest,
    ResumeDraftAIOptimizeResponse,
    ErrorResponse,
)
from app.auth import get_current_user
from app.services.claude_service import get_chat_response
from app.services.resume_pdf_service import generate_resume_pdf
from app.services.ats_scoring_service import compute_ats_score

router = APIRouter(prefix="/resume-drafts", tags=["resume-drafts"])

MAX_DRAFTS = 15
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc"}
MAX_UPLOAD_SIZE = 10 * 1024 * 1024  # 10 MB
PDF_MAGIC = b"%PDF"
DOCX_MAGIC = b"PK"


# ── Helpers ──────────────────────────────────────────────────────────────────

def _get_draft_or_404(draft_id: str, user_id: str, db: Session) -> ResumeDraft:
    draft = db.query(ResumeDraft).filter(
        ResumeDraft.id == draft_id,
        ResumeDraft.user_id == user_id,
    ).first()
    if not draft:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Resume draft not found",
        )
    return draft


def _validate_resume_json(raw: str) -> dict:
    """Parse and lightly validate resume JSON. Returns parsed dict."""
    try:
        data = json.loads(raw)
        if not isinstance(data, dict):
            raise ValueError
        return data
    except (json.JSONDecodeError, ValueError):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid resume JSON",
        )


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


# ── CRUD ─────────────────────────────────────────────────────────────────────

@router.get("", response_model=ResumeDraftListResponse)
def list_drafts(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    drafts = (
        db.query(ResumeDraft)
        .filter(ResumeDraft.user_id == current_user.id)
        .order_by(ResumeDraft.updated_at.desc())
        .all()
    )
    return ResumeDraftListResponse(
        drafts=[ResumeDraftResponse.model_validate(d) for d in drafts]
    )


@router.post(
    "",
    response_model=ResumeDraftResponse,
    status_code=status.HTTP_201_CREATED,
    responses={429: {"model": ErrorResponse}},
)
def create_draft(
    data: ResumeDraftCreateRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    count = db.query(ResumeDraft).filter(
        ResumeDraft.user_id == current_user.id,
    ).count()
    if count >= MAX_DRAFTS:
        raise HTTPException(
            status_code=status.HTTP_429_TOO_MANY_REQUESTS,
            detail=f"Maximum {MAX_DRAFTS} resumes allowed. Delete an existing one first.",
        )

    # Pre-fill with user profile data if available
    profile = db.query(UserProfile).filter(
        UserProfile.user_id == current_user.id
    ).first()
    initial_json = _build_initial_json(current_user, profile)

    draft = ResumeDraft(
        user_id=current_user.id,
        title=data.title or "Untitled Resume",
        template=data.template or "professional",
        source=data.source or "scratch",
        resume_json=json.dumps(initial_json),
    )
    db.add(draft)
    db.commit()
    db.refresh(draft)
    return ResumeDraftResponse.model_validate(draft)


@router.get("/{draft_id}", response_model=ResumeDraftResponse)
def get_draft(
    draft_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    draft = _get_draft_or_404(draft_id, current_user.id, db)
    return ResumeDraftResponse.model_validate(draft)


@router.delete("/{draft_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_draft(
    draft_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    draft = _get_draft_or_404(draft_id, current_user.id, db)
    db.delete(draft)
    db.commit()


# ── Autosave ─────────────────────────────────────────────────────────────────

@router.patch(
    "/{draft_id}",
    response_model=ResumeDraftResponse,
    responses={409: {"model": ErrorResponse}},
)
def autosave_draft(
    draft_id: str,
    data: ResumeDraftUpdateRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    draft = _get_draft_or_404(draft_id, current_user.id, db)

    # Optimistic locking: reject only if server is significantly newer (>5s)
    # Turso's eventual consistency can cause minor timestamp drift
    from datetime import datetime
    try:
        server_ts = datetime.fromisoformat(draft.updated_at.replace('Z', '+00:00'))
        client_ts = datetime.fromisoformat(data.updated_at.replace('Z', '+00:00'))
        diff = (server_ts - client_ts).total_seconds()
        if diff > 5:  # Only reject if >5 second gap (real conflict, not drift)
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail="This resume was edited elsewhere. Please reload.",
            )
    except (ValueError, TypeError):
        pass  # If timestamps can't be parsed, allow the save

    # Validate JSON is parseable
    _validate_resume_json(data.resume_json)

    draft.resume_json = data.resume_json
    draft.updated_at = _now_iso()
    db.commit()
    db.refresh(draft)
    return ResumeDraftResponse.model_validate(draft)


@router.patch(
    "/{draft_id}/template",
    response_model=ResumeDraftResponse,
)
def update_template(
    draft_id: str,
    data: dict,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    draft = _get_draft_or_404(draft_id, current_user.id, db)
    template = data.get("template", "professional")
    if template not in ("professional", "modern", "simple", "rendercv", "sidebar", "jake"):
        raise HTTPException(status_code=400, detail="Invalid template")
    draft.template = template
    draft.updated_at = _now_iso()
    db.commit()
    db.refresh(draft)
    return ResumeDraftResponse.model_validate(draft)


# ── Upload & Parse ───────────────────────────────────────────────────────────

@router.post(
    "/upload",
    response_model=ResumeDraftResponse,
    responses={400: {"model": ErrorResponse}, 429: {"model": ErrorResponse}},
)
async def upload_and_parse(
    file: UploadFile = File(...),
    template: str = Form(default="professional"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    # Check draft count
    count = db.query(ResumeDraft).filter(
        ResumeDraft.user_id == current_user.id,
    ).count()
    if count >= MAX_DRAFTS:
        raise HTTPException(
            status_code=status.HTTP_429_TOO_MANY_REQUESTS,
            detail=f"Maximum {MAX_DRAFTS} resumes allowed.",
        )

    # Validate file extension
    filename = file.filename or ""
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only PDF and DOCX files are supported.",
        )

    # Read file content with size limit
    content = await file.read()
    if len(content) > MAX_UPLOAD_SIZE:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File too large. Maximum size is 10MB.",
        )

    # Validate magic bytes
    if ext == ".pdf" and not content[:4].startswith(PDF_MAGIC):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="This file doesn't appear to be a valid PDF.",
        )
    if ext in (".docx", ".doc") and not content[:2].startswith(DOCX_MAGIC):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="This file doesn't appear to be a valid DOCX.",
        )

    # Extract text
    try:
        raw_text = _extract_text(content, ext)
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Could not read this file. It may be corrupted or password-protected.",
        )

    if len(raw_text.strip()) < 50:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Could not extract meaningful text. This may be a scanned/image-based PDF. Please upload a text-based file.",
        )

    # AI field mapping
    resume_json_str = await _parse_resume_with_ai(raw_text)

    # Validate template
    if template not in ("professional", "modern", "simple", "rendercv", "sidebar", "jake"):
        template = "professional"

    draft = ResumeDraft(
        user_id=current_user.id,
        title=f"Uploaded Resume",
        template=template,
        source="upload",
        resume_json=resume_json_str,
    )
    db.add(draft)
    db.commit()
    db.refresh(draft)
    return ResumeDraftResponse.model_validate(draft)


def _extract_text(content: bytes, ext: str) -> str:
    """Extract text from PDF or DOCX bytes."""
    if ext == ".pdf":
        return _extract_pdf_text(content)
    else:
        return _extract_docx_text(content)


def _extract_pdf_text(content: bytes) -> str:
    """Extract text from PDF using pdfplumber (already in requirements)."""
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for i, page in enumerate(pdf.pages[:5]):  # Max 5 pages
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def _extract_docx_text(content: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="DOCX parsing not available. Please upload a PDF instead.",
        )
    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also try to get text from tables
    for table in doc.tables:
        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells_text:
                paragraphs.append(" | ".join(cells_text))

    return "\n".join(paragraphs)


async def _parse_resume_with_ai(raw_text: str) -> str:
    """Use Claude to map raw resume text into structured JSON."""
    prompt = f"""You are a resume parser. Extract structured data from the following resume text.

Return ONLY valid JSON in this exact structure (no explanation, no markdown):
{{
    "personal_info": {{
        "name": "",
        "email": "",
        "phone": "",
        "location": "",
        "linkedin": null,
        "portfolio": null,
        "github": null
    }},
    "objective": "",
    "education": [
        {{
            "degree": "",
            "institution": "",
            "year": "",
            "grade": "",
            "board": null,
            "stream": ""
        }}
    ],
    "experience": [
        {{
            "title": "",
            "company": "",
            "duration": "",
            "location": "",
            "bullets": []
        }}
    ],
    "projects": [
        {{
            "name": "",
            "description": "",
            "tech_stack": [],
            "bullets": []
        }}
    ],
    "skills": {{
        "technical": [],
        "soft": [],
        "languages": [],
        "tools": []
    }},
    "achievements": [],
    "certifications": [
        {{
            "name": "",
            "issuer": "",
            "year": ""
        }}
    ]
}}

Rules:
- Use empty arrays [] for missing sections, never null
- Extract ALL information you can find
- For bullet points, use strong action verbs
- If a section is not found in the resume, leave its array/object empty
- Return ONLY the JSON, nothing else

Resume text:
{raw_text[:8000]}"""

    response = await get_chat_response(
        system_prompt="You are a precise resume parser. Output only valid JSON.",
        messages=[{"role": "user", "content": prompt}],
    )

    # Extract JSON from response (handle cases where AI wraps in markdown)
    text = response.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\n?", "", text)
        text = re.sub(r"\n?```$", "", text)

    # Validate JSON
    try:
        data = json.loads(text)
        if not isinstance(data, dict):
            raise ValueError
        return json.dumps(data)
    except (json.JSONDecodeError, ValueError):
        # Return empty structure if AI fails
        return json.dumps(_empty_resume_json())


# ── AI Per-Field Optimization ────────────────────────────────────────────────

@router.post(
    "/{draft_id}/ai-optimize",
    response_model=ResumeDraftAIOptimizeResponse,
)
async def ai_optimize_field(
    draft_id: str,
    data: ResumeDraftAIOptimizeRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _get_draft_or_404(draft_id, current_user.id, db)

    if not data.current_value.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Field is empty. Type something first, then optimize.",
        )

    context_str = ""
    if data.context:
        context_str = f"\nContext about the person: {json.dumps(data.context)}"

    field_prompts = {
        "objective": "Rewrite this professional summary to be concise (2-3 sentences), ATS-friendly, and use strong keywords relevant to their field. Keep it specific, not generic.",
        "experience_bullet": "Rewrite this work experience bullet point to start with a strong action verb, include quantifiable results where possible, and be ATS-keyword-rich. Keep it to one sentence.",
        "project_description": "Rewrite this project description to sound more technical and impactful. Mention technologies used and the outcome/impact.",
        "project_bullet": "Rewrite this project bullet point to start with an action verb and highlight technical skills and measurable outcomes.",
        "skills": "Based on the listed skills and context, suggest 3-5 additional relevant technical skills that are commonly expected for this kind of profile. Return only the skill names separated by commas.",
    }

    field_type = data.field
    instruction = field_prompts.get(field_type, field_prompts["experience_bullet"])

    prompt = f"""{instruction}
{context_str}

Current text: "{data.current_value}"

Return ONLY the improved text, nothing else. No quotes, no explanation."""

    response = await get_chat_response(
        system_prompt="You are an expert resume writer and ATS optimization specialist. You write concise, impactful, keyword-rich resume content.",
        messages=[{"role": "user", "content": prompt}],
    )

    return ResumeDraftAIOptimizeResponse(optimized_value=response.strip())


# ── ATS Score ────────────────────────────────────────────────────────────────

@router.post("/{draft_id}/ats-score")
async def get_ats_score(
    draft_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    draft = _get_draft_or_404(draft_id, current_user.id, db)
    resume_data = _validate_resume_json(draft.resume_json)
    result = await compute_ats_score(resume_data)

    # Cache the total score
    if isinstance(result, dict) and "total_score" in result:
        draft.ats_score = result["total_score"]
        db.commit()

    return result


# ── PDF Download ─────────────────────────────────────────────────────────────

@router.get("/{draft_id}/download")
def download_pdf(
    draft_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    draft = _get_draft_or_404(draft_id, current_user.id, db)

    try:
        resume_data = json.loads(draft.resume_json)
    except json.JSONDecodeError:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Resume data is invalid.",
        )

    # Sanitize: replace None values with empty strings to prevent reportlab crashes
    _sanitize_resume_data(resume_data)

    # Get profile image URL if available
    profile = db.query(UserProfile).filter(
        UserProfile.user_id == current_user.id
    ).first()
    profile_image_url = getattr(profile, "profile_image", None) if profile else None

    try:
        pdf_bytes = generate_resume_pdf(
            resume_json=resume_data,
            template=draft.template,
            profile_image_url=profile_image_url,
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"PDF generation failed: {type(e).__name__}: {str(e)[:200]}",
        )

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="resume-{draft.template}.pdf"',
        },
    )


# ── Helpers ──────────────────────────────────────────────────────────────────

def _sanitize_resume_data(data: dict) -> None:
    """Replace None with '' in resume data to prevent reportlab crashes."""
    if not isinstance(data, dict):
        return
    for key, val in data.items():
        if val is None:
            data[key] = ""
        elif isinstance(val, dict):
            _sanitize_resume_data(val)
        elif isinstance(val, list):
            for i, item in enumerate(val):
                if item is None:
                    val[i] = ""
                elif isinstance(item, dict):
                    _sanitize_resume_data(item)
                elif not isinstance(item, str):
                    val[i] = str(item)


def _build_initial_json(user, profile=None) -> dict:
    """Build initial resume JSON pre-filled with user profile data."""
    data = _empty_resume_json()
    data["personal_info"]["name"] = user.name or ""

    if profile:
        if profile.city and profile.state:
            data["personal_info"]["location"] = f"{profile.city}, {profile.state}"
        if profile.education_level and profile.stream:
            data["education"] = [{
                "degree": f"{profile.education_level} {profile.stream}".strip(),
                "institution": user.college or "",
                "year": profile.class_or_year or "",
                "grade": f"CGPA: {profile.cgpa}" if profile.cgpa else "",
                "board": profile.board,
                "stream": profile.stream or "",
            }]
        elif user.college:
            data["education"] = [{
                "degree": "",
                "institution": user.college,
                "year": "",
                "grade": "",
                "board": None,
                "stream": "",
            }]

        # Pre-fill skills from profile
        try:
            skills = json.loads(profile.skills) if profile.skills else []
            if isinstance(skills, list):
                data["skills"]["technical"] = skills
        except (json.JSONDecodeError, TypeError):
            pass

        if profile.career_aspiration_raw:
            data["objective"] = profile.career_aspiration_raw

    return data


def _empty_resume_json() -> dict:
    return {
        "personal_info": {
            "name": "",
            "email": "",
            "phone": "",
            "location": "",
            "linkedin": None,
            "portfolio": None,
            "github": None,
        },
        "objective": "",
        "education": [],
        "experience": [],
        "projects": [],
        "skills": {
            "technical": [],
            "soft": [],
            "languages": [],
            "tools": [],
        },
        "achievements": [],
        "certifications": [],
    }
